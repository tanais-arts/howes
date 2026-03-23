from docx import Document
from docx.shared import Pt
import tempfile
import os

def create_report(chapters, template_path=None):
    """
    chapters: list of dict {title, start, end, first: [w...], last: [w...]}
    Returns path to generated .docx
    """
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    doc.add_heading('Extrait des chapitres — 3 premiers / 3 derniers mots', level=1)
    for idx, c in enumerate(chapters, 1):
        title = c.get('title') or f'Chapitre {idx}'
        doc.add_heading(title, level=2)
        doc.add_paragraph(f"Pages: {c.get('start')+1}–{c.get('end')+1}")
        first = ' '.join(c.get('first', []))
        last = ' '.join(c.get('last', []))
        doc.add_paragraph(f"Premiers mots: {first}")
        doc.add_paragraph(f"Derniers mots: {last}")

    fd, out = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    doc.save(out)
    return out
